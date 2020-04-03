from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import RGBColor
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches
#信息 
namelist = ['赵云鹏','蔡敏琪','陈瑶','杨余彬']
collect_num = 'HZU00251'
copy_num = 1
date = '2015-08-03'
Family = 'Verbenaceae'
LocalName = '紫珠属'
Species = 'Callicarpa'
Location = '中国浙江省临安市西天目山坞子岭二号样地'
Lon = '119°26\'19.752\"E'
Lat = '30°18\'48.942\"N'
Frequency = '偶见'
Alt = '352'
Habitat = '(马尾松)针阔混交林'
Substrate = ''
Habit = ''
Life = ''
Height = ''
Diam = ''
Root = ''
Stem = ''
Leaf = ''
Flower = ''
Fruit = ''
Seed = ''
Introduction = ''
MolecularSample = ''
Additional = ''

#function
alpha = ' '.join([chr(x) for x in range(32,127)])
def mystyle1(argc1,str1):# 1attribute per line
    paragraph = document.add_paragraph()
    paragraph.add_run(str1)
    run = paragraph.add_run(set_align(argc1,cal(str1+argc1)))
    run.bold = True
    run.underline = True


def mystyle2(argc1,argc2,str1,str2): # 2 attribute per line
    paragraph = document.add_paragraph()
    paragraph.add_run(str1)
    run = paragraph.add_run('______{0}______'.format(argc1))
    run.bold = True
    run.underline = True
    paragraph.add_run(str2)
    run = paragraph.add_run(set_align('______{0}______'.format(argc2),cal(str1+('______{0}______'.format(argc1))+str2+('______{0}______'.format(argc2)))))
    run.bold = True
    run.underline = True

def cal(s):
    alpha_num = len([x for x in s if x in alpha])
    hanzi_num = len(s) - alpha_num
    tot = (hanzi_num << 1) + alpha_num
    ans = 75 - tot
    return ans

def set_align(ret,tot):
    if tot & 1:
        ret += '_'
        tot -= 1
    tot >>= 1
    ret = '_'*tot + ret + '_'*tot
    return ret

#处理
#标题
document = Document()
document.styles['Normal'].font.name = u'宋体'
document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
paragraph = document.add_paragraph()
paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
paragraph.add_run('华东山地生物多样性调查').bold = True

#内容
mystyle1(','.join(namelist),'采集人 Collector(s):')
mystyle1(collect_num,'采集号 No.:')
mystyle2(copy_num,date,'标本份数 Copy:','日期 Date:')
mystyle1(Family,'科名 Family:')
mystyle1(LocalName,'俗名 Local Name:')
mystyle1(Species,'学名 Species:')
mystyle1(Location,'产地 Location:')
mystyle2(Lon,Lat,'经度 Lon.','纬度 Lat.')
mystyle2(Frequency,Alt,'丰富度 Frequency:','海拔 Alt:')
mystyle1(Habitat,'生境 Habitat:')
mystyle1(Substrate,'基质 Substrate:')
mystyle2(Habit,Life,'性状 Habit:','生活型 Life:')
mystyle2(Root,Stem,'根 Root:','径 Stem:')
mystyle2(Leaf,Flower,'叶 Leaf:','花 Flower:')
mystyle2(Fruit,Seed,'果实 Fruit:','种子 Seed:')
mystyle1(Introduction,'引种 Introduction:')
mystyle1(MolecularSample,'分子材料 MolecularSample:')
mystyle1(Additional,'附记 Additional:')
document.save('test1.docx')
